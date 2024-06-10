from typing import List, Dict
import uuid
from datetime import datetime
import os

from django.shortcuts import get_object_or_404
from django.http import FileResponse
from ninja import Router, Query, UploadedFile, File, Form
from ninja.pagination import paginate

from .schemas.solicitacao import (
    SolicitacaoIn, SolicitacaoOut, FilterSolicitacao, SolicitacaoUpdate)
from .schemas.ponto_coleta import PontoColetaIn, PontoColetaOut
from .. import models
from .utils import save_file

from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn, nsdecls
from docx.oxml import OxmlElement, parse_xml
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from io import BytesIO

router = Router(tags=["Solicitacoes"])


@router.get("/", response=List[SolicitacaoOut])
@paginate
def list_solicitacoes(request, filters: FilterSolicitacao = Query(...)):
    qs = models.Solicitacao.objects.all()
    return filters.filter(qs)


@router.get("/{id}", response=SolicitacaoOut)
def get_solicitacao(request, id: int):
    qs = get_object_or_404(models.Solicitacao, id=id)
    print(qs)
    return qs


@router.post("/{id}/imagem")
def upload_image(request, id: int, description: str = Form(...), file: UploadedFile = File(...)):
    solicitacao = get_object_or_404(models.Solicitacao, id=id)

    img_path = save_file(f'media/images/solicitacoes/solicitacao_{solicitacao.id}_{uuid.uuid4()}.png', file)
    image = models.Image.objects.create(file=img_path, description=description)
    image.save()

    solicitacao.imagens.add(image)
    solicitacao.save()

    return {"success": True}


@router.post("/", response=SolicitacaoOut)
def create_solicitacao(request, payload: SolicitacaoIn):
    data = payload.dict()
    ponto = get_object_or_404(models.PontoColeta, id=data.pop("ponto_id"))
    data["ponto"] = ponto
    solicitacao = models.Solicitacao.objects.create(**data)
    solicitacao.save()
    return solicitacao

@router.put("/{id}", response=SolicitacaoOut)
def update_solicitacao(request, id: int, payload: SolicitacaoUpdate):
    data = payload.dict()

    solicitacao = get_object_or_404(models.Solicitacao, id=id)
    ponto = get_object_or_404(models.PontoColeta, id=data.pop("ponto_id"))
    data["ponto"] = ponto

    for attr, value in data.items():
        setattr(solicitacao, attr, value)
    solicitacao.save()
    return solicitacao


@router.delete("/{id}")
def delete_solicitacao(request, id: int):
    solicitacao = get_object_or_404(models.Solicitacao, id=id)
    solicitacao.delete()
    return {"success": True}

@router.get("/{id}/document")
def generate_doc(request, id: int):
    data = get_object_or_404(models.Solicitacao, id=id)
    if not data:
        return FileResponse("Solicitação não encontrada", status=404)

    doc = Document()

    if data.tipo == 'Limpeza de reservatório':
        doc.add_heading(
            'Solicitação de limpeza de reservatório predial de água na UFERSA campus Mossoró', level=1)

    # Table with 2 columns
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Table Grid'

    # Row Title
    row = table.rows[0]
    cell = row.cells[0]
    cell.text = "Descrição da Solicitação"
    cell.merge(row.cells[1])

    # Centralize o texto
    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Row 1
    cell = table.cell(1, 0)
    cell.text = 'Edificação'
    cell = table.cell(1, 1)
    campus = "Leste" if "LE" in data.ponto.edificacao.campus else "Oeste"
    cell.text = data.ponto.edificacao.nome + ', lado ' + campus

    # Row 2
    cell = table.cell(2, 0)
    cell.text = 'Serviço solicitado'
    cell = table.cell(2, 1)
    cell.text = data.tipo

    # Row 3
    cell = table.cell(3, 0)
    cell.text = 'Objetivo do serviço solicitado'
    paragraph = cell.paragraphs[0]
    vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table.cell(3, 0).vertical_alignment = vertical_alignment

    cell = table.cell(3, 1)
    cell.text = data.objetivo

    # Row 4 (merged cells)
    cell = table.cell(4, 0)
    cell.text = "Justificativa da Solicitação"
    paragraph = cell.paragraphs[0]

    vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table.cell(4, 0).vertical_alignment = vertical_alignment

    cell = table.cell(4, 1)

    cell.text = data.justificativa
    cell = table.cell(5, 0)

    cell.text = "Nº da solicitação/ano"
    cell = table.cell(5, 1)

    data_str = data.data.strftime("%m/%Y")
    cell.text = data_str


    # Footnotes
    if data.tipo == "Conserto de reservatório":
        doc.add_paragraph(
            '\n¹ ASSOCIAÇÃO BRASILEIRA DE NORMAS TÉCNICAS (ABNT). NBR 5626: sistemas prediais de água fria e água quente – projeto, execução, operação e manutenção. S.l.: ABNT, 2020. 55 p.')
        doc.add_paragraph('² BRASIL. Ministério da Saúde. Portaria de Consolidação GM/MS nº 5, de 28 de setembro de 2017. Altera o Anexo XX da Portaria de Consolidação GM/MS nº 5, de 28 de setembro de 2017, para dispor sobre os procedimentos de controle e de vigilância da qualidade da água para consumo humano e seu padrão de potabilidade. 2017. Disponível em: https://bvsms.saude.gov.br/bvs/saudelegis/gm/2017/prc0005_30_10_2017.html#ANEXOXX. Acesso em: 12 fev. 2023.')
        doc.add_paragraph('³ BRASIL. Ministério da Saúde. Portaria GM/MS nº 888, de 4 de maio de 2021. Altera o Anexo XX da Portaria de Consolidação GM/MS nº 5, de 28 de setembro de 2017, para dispor sobre os procedimentos de controle e de vigilância da qualidade da água para consumo humano e seu padrão de potabilidade. 2021a. Disponível em: https://www.in.gov.br/en/web/dou/-/portaria-gm/ms-n-888-de-4-de-maio-de-2021-321540185. Acesso em: 12 fev. 2023.')

    elif data.tipo == "Limpeza de reservatório":
        doc.add_paragraph(
            '\n¹ ASSOCIAÇÃO BRASILEIRA DE NORMAS TÉCNICAS (ABNT). NBR 5626: sistemas prediais de água fria e água quente – projeto, execução, operação e manutenção. S.l.: ABNT, 2020. 55 p.')


    # Se houver imagens, adicione uma nova página
    image_paths = data.imagens.all().values()
    if image_paths:
        doc.add_page_break()
        doc.add_heading('Vistas do local', level=1)

        # Table for images and captions
        table = doc.add_table(rows=0, cols=1)
        table.style = 'Table Grid'


        # Add images and captions to table
        for idx, image in enumerate(image_paths):
            # Adicione uma nova linha para a imagem
            table.add_row()
            row_cells = table.rows[-1].cells
            p = row_cells[0].paragraphs[0]
            run = p.add_run()
            # Set the width to 2 inches
            run.add_picture("src/files/" + image["file"], width=Inches(2))
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Alinhamento horizontal

            # Adicione uma nova linha para a descrição
            table.add_row()
            row_cells = table.rows[-1].cells
            p = row_cells[0].add_paragraph()
            p.add_run(image["description"])
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Ajuste o alinhamento vertical das células que contêm as imagens
        for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Save the document

    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return FileResponse(doc_io, as_attachment=True, filename="solicitacao.docx")
