import os
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL

from docx.shared import Inches

from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

from datetime import datetime

data = {
  "id": 1,
  "ponto": {
    "id": 26,
    "imagem": None,
    "ambiente": "Hall de entrada",
    "tipo": 1,
    "tombo": "-",
    "edificacao": {
      "imagem": None,
      "codigo": "A11",
      "nome": "Prédio Central",
      "campus": "OE",
      "cronograma": 6,
      "pontos_url": "/api/v1/edificacoes/A11/pontos",
      "imagens": []
    },
    "edificacao_url": "/api/v1/edificacoes/A11",
    "fluxos_url": "/api/v1/pontos/26/fluxos",
    "status": None,
    "status_message": "Não há coletas nesse ponto",
    "amontante": None,
    "associados": []
  },
  "data": "2024-05-27T02:14:55.820Z",
  "tipo": "Limpeza de reservatório",
  "objetivo": "Contribuir para a preservação da potabilidade da água para consumo humano da UFERSA.",
  "justificativa": "- Comprovação da necessidade de limpeza do reservatório de água, a partir de amostragem de água no âmbito do projeto “Qualidade da água para consumo humano: estudo no sistema da UFERSA-Mossoró” (cadastro na PROPPG: PIB10009-2019).\n\n- Preservação da potabilidade da água conforme previsto na NBR 5626/2020 (ABNT, 2020, p. 40)¹:\n[...] “Todas as partes acessíveis dos componentes que têm contato com a água devem ser limpas periodicamente.” [...]\n\nObs.: Para limpeza de reservatório de água, recomenda-se o procedimento especificado no Anexo F da NBR 5626/2020.",
  "imagens": [
    {
      "file": "/files/media/images/solicitacoes/solicitacao_1_e3854306-26a3-4ab1-af21-e4df8db97c1b.png",
      "description": "teste"
    },
    {
      "file": "/files/media/images/solicitacoes/solicitacao_1_e3854306-26a3-4ab1-af21-e4df8db97c1b.png",
      "description": "teste"
    },
    {
      "file": "/files/media/images/solicitacoes/solicitacao_1_e3854306-26a3-4ab1-af21-e4df8db97c1b.png",
      "description": "teste"
    }
  ],
  "status": "Pendente"
}

file_path = 'src/files/media/images/solicitacoes/solicitacao_1_e3854306-26a3-4ab1-af21-e4df8db97c1b.png'
if os.path.isfile(file_path):
    print("File exists")
else:
    print("File does not exist")


# Create a new Document
doc = Document()

# Title
if data['tipo'] == 'Limpeza de reservatório':
    doc.add_heading('Solicitação de limpeza de reservatório predial de água na UFERSA campus Mossoró', level=1)

# Table with 2 columns
table = doc.add_table(rows=6, cols=2)
table.style = 'Table Grid'

# Row Title
row = table.rows[0]
cell = row.cells[0]
cell.text = "Descrição da Solicitação"
cell.merge(row.cells[1])

# Centralize o texto
cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Row 1
cell = table.cell(1, 0)
cell.text = 'Edificação'
cell = table.cell(1, 1)
campus = "Leste" if "LE" in data["ponto"]["edificacao"]["campus"] else "Oeste"
cell.text = data['ponto']['edificacao']['nome'] + ', lado ' + campus

# Row 2
cell = table.cell(2, 0)
cell.text = 'Serviço solicitado'
cell = table.cell(2, 1)
cell.text = data["tipo"]

# Row 3
cell = table.cell(3, 0)
cell.text = 'Objetivo do serviço solicitado'
paragraph = cell.paragraphs[0]
vertical_alignment = WD_ALIGN_VERTICAL.CENTER
table.cell(3, 0).vertical_alignment = vertical_alignment

cell = table.cell(3, 1)
cell.text = data["objetivo"]

# Row 4 (merged cells)
cell = table.cell(4, 0)
cell.text = "Justificativa da Solicitação"
paragraph = cell.paragraphs[0]
# paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
vertical_alignment = WD_ALIGN_VERTICAL.CENTER
table.cell(4, 0).vertical_alignment = vertical_alignment

# cell.merge(table.cell(3, 1))
cell = table.cell(4, 1)
cell.text = data["justificativa"]

# Next part
cell = table.cell(5, 0)
cell.text = "Nº da solicitação/ano"

cell = table.cell(5, 1)

# Convertendo a string para um objeto datetime
date_object = datetime.strptime(data["data"], "%Y-%m-%dT%H:%M:%S.%fZ")

# Formatando o objeto datetime para mostrar apenas o mês e o ano
formatted_date = date_object.strftime("%m/%Y")
cell.text = formatted_date
# doc.add_paragraph(data["data"])
# doc.add_paragraph('13/2023')

# Footnotes
if data["tipo"] == "Conserto de reservatório":
    doc.add_paragraph('\n¹ ASSOCIAÇÃO BRASILEIRA DE NORMAS TÉCNICAS (ABNT). NBR 5626: sistemas prediais de água fria e água quente – projeto, execução, operação e manutenção. S.l.: ABNT, 2020. 55 p.')
    doc.add_paragraph('² BRASIL. Ministério da Saúde. Portaria de Consolidação GM/MS nº 5, de 28 de setembro de 2017. Altera o Anexo XX da Portaria de Consolidação GM/MS nº 5, de 28 de setembro de 2017, para dispor sobre os procedimentos de controle e de vigilância da qualidade da água para consumo humano e seu padrão de potabilidade. 2017. Disponível em: https://bvsms.saude.gov.br/bvs/saudelegis/gm/2017/prc0005_30_10_2017.html#ANEXOXX. Acesso em: 12 fev. 2023.')
    doc.add_paragraph('³ BRASIL. Ministério da Saúde. Portaria GM/MS nº 888, de 4 de maio de 2021. Altera o Anexo XX da Portaria de Consolidação GM/MS nº 5, de 28 de setembro de 2017, para dispor sobre os procedimentos de controle e de vigilância da qualidade da água para consumo humano e seu padrão de potabilidade. 2021a. Disponível em: https://www.in.gov.br/en/web/dou/-/portaria-gm/ms-n-888-de-4-de-maio-de-2021-321540185. Acesso em: 12 fev. 2023.')

elif data["tipo"] == "Limpeza de reservatório":
    doc.add_paragraph('\n¹ ASSOCIAÇÃO BRASILEIRA DE NORMAS TÉCNICAS (ABNT). NBR 5626: sistemas prediais de água fria e água quente – projeto, execução, operação e manutenção. S.l.: ABNT, 2020. 55 p.')


# New page for images and descriptions
doc.add_page_break()
doc.add_heading('Vistas do local', level=1)

# Table for images and captions
table = doc.add_table(rows=0, cols=1)
table.style = 'Table Grid'

# Adding images and their captions

# Note the parentheses around the image path and caption, making them a tuple
# image_paths = [("src/files/media/images/solicitacoes/solicitacao_1_e3854306-26a3-4ab1-af21-e4df8db97c1b.png",
#                 "Fotografia em 16/03/2023: vista da fachada frontal do Laboratório de Engenharias II"),
#                 ("src/files/media/images/solicitacoes/solicitacao_1_e3854306-26a3-4ab1-af21-e4df8db97c1b.png",
#                 "Fotografia em 16/03/2023: vista da fachada frontal do Laboratório de Engenharias II"),
#                 ("src/files/media/images/solicitacoes/solicitacao_1_e3854306-26a3-4ab1-af21-e4df8db97c1b.png",
#                 "Fotografia em 16/03/2023: vista da fachada frontal do Laboratório de Engenharias II"),
#                 ("src/files/media/images/solicitacoes/solicitacao_1_e3854306-26a3-4ab1-af21-e4df8db97c1b.png",
#                 "Fotografia em 16/03/2023: vista da fachada frontal do Laboratório de Engenharias II")]

image_paths = data["imagens"]


# Add images and captions to table
for idx, image in enumerate(image_paths):
    # Adicione uma nova linha para a imagem
    table.add_row()
    row_cells = table.rows[-1].cells
    p = row_cells[0].paragraphs[0]
    run = p.add_run()
    # Set the width to 2 inches
    run.add_picture("src/" + image["file"], width=Inches(2))
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Alinhamento horizontal

    # Adicione uma nova linha para a descrição
    table.add_row()
    row_cells = table.rows[-1].cells
    p = row_cells[0].add_paragraph()
    p.add_run(image["description"])
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Ajuste o alinhamento vertical das células que contêm as imagens
for row in table.rows:
    for cell in row.cells:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

# Save the document
doc.save("soliciacao.docx")
